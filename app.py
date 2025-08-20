import streamlit as st
import requests
import uuid
import os
import io
import asyncio
import edge_tts
import re
from docx import Document
from reportlab.pdfgen import canvas

# =========================
# Paths & Config
# =========================
PAPER_FOLDER = r"C:\Users\mohamed mahmoud emam\OneDrive\Desktop\Research Paper Summarizer\ResearchPapers"
os.makedirs(PAPER_FOLDER, exist_ok=True)

# Page Config
st.set_page_config(
    page_title="📚 Research Paper Summarizer & Chatbot",
    page_icon="📖",
    layout="wide"
)

# =========================
# Header
# =========================
st.markdown(
    """
    <div style="text-align:center; padding: 10px;">
        <h1>📚 Research Paper Summarizer & Chatbot</h1>
        <p style="color: #555; font-size: 18px;">
            Upload or fetch research papers, get summaries, translations, and chat with your documents.
        </p>
    </div>
    """,
    unsafe_allow_html=True
)
st.divider()

# =========================
# Session State Init
# =========================
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())

for key, default in {
    "chat_history": [],
    "summary": None,
    "translated_summary": None,
    "tts_audio": None,
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

n_results = 5  # fixed

# =========================
# Upload & Fetch Section
# =========================
st.subheader("📂 Add Your Research Paper")

with st.container():
    col1, col2 = st.columns(2)

    # --- Upload PDF ---
    with col1:
        st.markdown("#### ⬆️ Upload a PDF")
        uploaded_file = st.file_uploader("Choose a research paper", type=["pdf"])
        if uploaded_file and st.button("📤 Upload", use_container_width=True):
            save_path = os.path.join(PAPER_FOLDER, uploaded_file.name)
            try:
                with open(save_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
                st.success(f"✅ Saved to {save_path}")

                with open(save_path, "rb") as f:
                    response = requests.post("http://127.0.0.1:8000/upload", files={"file": f})

                if response.status_code == 200:
                    st.success("📤 Uploaded & Indexed Successfully!")
                else:
                    st.error("❌ Backend Upload Failed.")
            except Exception as e:
                st.error(f"⚠️ Backend Connection Error: {e}")

    # --- Fetch from Link ---
    with col2:
        st.markdown("#### 🔗 Fetch from a Link")
        paper_link = st.text_input("Paste a paper link (PDF / Webpage)")
        if st.button("🌐 Fetch & Index", use_container_width=True):
            if paper_link:
                with st.spinner("Fetching paper... ⏳"):
                    try:
                        response = requests.post("http://127.0.0.1:8000/fetch_link", json={"url": paper_link})
                        if response.status_code == 200:
                            st.success("✅ Paper Fetched & Indexed Successfully!")
                        else:
                            st.error("❌ Could not fetch the paper.")
                    except Exception as e:
                        st.error(f"⚠️ Backend Connection Error: {e}")
            else:
                st.warning("⚠️ Please enter a valid link.")

st.divider()

# =========================
# Summarization
# =========================
st.subheader("📝 Summarize Your Papers")

if st.button("📊 Generate Summary", use_container_width=True):
    with st.spinner("Summarizing... ⏳"):
        try:
            response = requests.get("http://127.0.0.1:8000/summarize")
            if response.status_code == 200:
                st.session_state.summary = response.json().get("summary", "No summary available.")
                st.success("✅ Summary Generated!")
            else:
                st.error("❌ API Error while summarizing.")
        except Exception as e:
            st.error(f"⚠️ Could not connect to API: {e}")

# =========================
# Display Summary & Translate
# =========================
if st.session_state.summary:
    summary = st.session_state.summary
    st.markdown("### 📄 Summary")
    st.info(summary)

    with st.expander("🌍 Translate Summary", expanded=False):
        target_lang = st.selectbox("Choose target language:", ["Arabic", "French", "Spanish", "German"])
        if st.button("Translate", use_container_width=True):
            with st.spinner("Translating... ⏳"):
                try:
                    response = requests.post(
                        "http://127.0.0.1:8000/translate",
                        json={"text": summary, "target_lang": target_lang}
                    )
                    if response.status_code == 200:
                        st.session_state.translated_summary = response.json().get("translated_text", "No translation available.")
                        st.success("✅ Translation Completed!")
                    else:
                        st.error("❌ API Error while translating.")
                except Exception as e:
                    st.error(f"⚠️ Could not connect to API: {e}")

        if st.session_state.translated_summary:
            st.markdown("#### 📄 Translated Summary")
            st.success(st.session_state.translated_summary)

    # --- Downloads ---
    with st.expander("💾 Download Options", expanded=False):
        colA, colB, colC = st.columns(3)

        with colA:
            st.download_button("⬇️ TXT", summary, "summary.txt", "text/plain")

        with colB:
            docx_buffer = io.BytesIO()
            doc = Document()
            doc.add_heading("Research Paper Summary", 0)
            doc.add_paragraph(summary)
            doc.save(docx_buffer)
            st.download_button(
                "⬇️ DOCX",
                docx_buffer.getvalue(),
                "summary.docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        with colC:
            pdf_buffer = io.BytesIO()
            pdf = canvas.Canvas(pdf_buffer)
            text_object = pdf.beginText(40, 800)
            text_object.setFont("Helvetica", 12)
            for line in summary.split("\n"):
                text_object.textLine(line)
            pdf.drawText(text_object)
            pdf.save()
            st.download_button(
                "⬇️ PDF",
                pdf_buffer.getvalue(),
                "summary.pdf",
                "application/pdf"
            )

    # --- TTS ---
    with st.expander("🔊 Listen to Summary", expanded=False):

        def clean_text(text):
            text = re.sub(r'<.*?>', '', text)
            text = re.sub(r'#\w+', '', text)
            text = re.sub(r'[^\w\s.,!?]', '', text)
            return text.strip()

        if st.button("▶️ Generate Audio", use_container_width=True):
            cleaned_summary = clean_text(summary)

            async def generate_speech():
                tts_file = "summary_audio.mp3"
                communicate = edge_tts.Communicate(cleaned_summary, voice="en-US-AriaNeural")
                await communicate.save(tts_file)
                with open(tts_file, "rb") as f:
                    return f.read()

            st.session_state.tts_audio = asyncio.run(generate_speech())
            st.success("✅ Audio Generated!")

        if st.session_state.tts_audio:
            st.audio(st.session_state.tts_audio, format="audio/mp3")
            st.download_button("⬇️ Download MP3", st.session_state.tts_audio, "summary.mp3", "audio/mp3")

st.divider()

# =========================
# Chat Section
# =========================
st.subheader("💬 Chat with Your Research Papers")

chat_container = st.container()
with chat_container:
    for msg in st.session_state.chat_history:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

if question := st.chat_input("Ask a question..."):
    st.session_state.chat_history.append({"role": "user", "content": question})

    with st.chat_message("user"):
        st.markdown(question)

    with st.chat_message("assistant"):
        with st.spinner("Retrieving answer... ⏳"):
            try:
                response = requests.post(
                    "http://127.0.0.1:8000/ask",
                    json={"question": question, "n_results": n_results}
                )
                if response.status_code == 200:
                    answer = response.json().get("answer", "No answer available.")
                    st.markdown(answer)
                    st.session_state.chat_history.append({"role": "assistant", "content": answer})
                else:
                    st.error("❌ API Error while answering.")
            except Exception as e:
                st.error(f"⚠️ Could not connect to API: {e}")